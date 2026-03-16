from __future__ import annotations

from pathlib import Path
from typing import List, Tuple

from openpyxl import load_workbook

from app.schemas.options import AnonymizeOptions
from app.services.anonymizer import (
    anonymize_text_value,
    apply_plain_placeholders,
    detect_spans_with_llm,
    replace_spans,
)

try:
    from docx import Document  # type: ignore
except Exception:
    Document = None  # type: ignore


def _iter_docx_paragraphs(doc) -> List:
    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for section in getattr(doc, "sections", []):
        paragraphs.extend(section.header.paragraphs)
        paragraphs.extend(section.footer.paragraphs)
    return paragraphs


def process_docx_file(src: Path, dst: Path, options: AnonymizeOptions) -> Tuple[str, List[str]]:
    if Document is None:
        raise RuntimeError("python-docx не установлен")
    doc = Document(str(src))
    preview_parts: List[str] = []

    for paragraph in _iter_docx_paragraphs(doc):
        if not paragraph.runs:
            continue
        full_text = paragraph.text
        if not full_text:
            continue
        spans = detect_spans_with_llm(full_text, options)
        spans = apply_plain_placeholders(spans)
        if not spans:
            if len(preview_parts) < 20:
                preview_parts.append(full_text)
            continue
        _apply_spans_to_runs(paragraph.runs, spans)
        if len(preview_parts) < 20:
            preview_parts.append(replace_spans(full_text, spans))

    doc.save(str(dst))
    return "\n".join(preview_parts[:20]), [
        "DOCX обработан с сохранением структуры документа. Если персональные данные разбиты по нескольким run-элементам, часть совпадений может потребовать доработки правил."
    ]


def process_xlsx_file(src: Path, dst: Path, options: AnonymizeOptions) -> Tuple[str, List[str]]:
    wb = load_workbook(str(src))
    preview_lines: List[str] = []

    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and cell.value:
                    cell.value = anonymize_text_value(cell.value, options, prefer_llm=True)
                    if len(preview_lines) < 20:
                        preview_lines.append(f"[{ws.title}] {cell.coordinate}: {cell.value}")

    wb.save(str(dst))
    return "\n".join(preview_lines), []


def _apply_spans_to_runs(runs, spans) -> None:
    # Build run offsets and buffers based on original text
    offsets = []
    buffers = []
    pos = 0
    for run in runs:
        text = run.text or ""
        start = pos
        end = pos + len(text)
        offsets.append((start, end))
        buffers.append(list(text))
        pos = end

    for span in sorted(spans, key=lambda s: s.start, reverse=True):
        first_replaced = False
        for idx, run in enumerate(runs):
            run_start, run_end = offsets[idx]
            if run_end <= span.start or run_start >= span.end:
                continue
            overlap_start = max(span.start, run_start) - run_start
            overlap_end = min(span.end, run_end) - run_start
            if overlap_start >= overlap_end:
                continue
            buf = buffers[idx]
            if not first_replaced:
                buf[overlap_start] = span.label
                for i in range(overlap_start + 1, overlap_end):
                    buf[i] = ""
                first_replaced = True
            else:
                for i in range(overlap_start, overlap_end):
                    buf[i] = ""

    for run, buf in zip(runs, buffers):
        run.text = "".join(ch for ch in buf if ch)
